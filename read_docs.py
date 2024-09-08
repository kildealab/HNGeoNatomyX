import textract
from docx import Document
import os
import subprocess
import aspose.words as aw
import matplotlib.image as mpimg

import textract
import PyPDF2
from PyPDF2 import PdfReader


def generate_docx(doc_path, path):

    subprocess.call(['soffice',
                 # '--headless',
                 '--convert-to',
                 'docx',
                 '--outdir',
                 path,
                 doc_path])
    return doc_path


def generate_pdf(doc_path, path):

    subprocess.call(['soffice',
                 # '--headless',
                 '--convert-to',
                 'pdf',
                 '--outdir',
                 path,
                 doc_path])
    return doc_path


#FUNCTION TO READ THE HABITS OR THE SOCIAL HISTORY OF THE PATIENT, TO EXTRACT RELEVANT INFORMATION
#SUCH AS IF THE PATIENT IS A SMOKER OR NOT, ALCOHOLIC, ETC.
def read_docs_habits_and_social_history(dir_list,PATH_DOCS):
    data_habits = []
    not_saved= []
  
    for j in dir_list:
      if j.split('.')[1]!='pdf':
          if j.split('.')[1]!='doc':
             doc = aw.Document(PATH_DOCS+j)
             formFields = doc.range.form_fields
             try:
               for p in doc.first_section.body.tables[1]:
                  row = p.as_row()
                  text = row.first_cell.to_string(aw.SaveFormat.TEXT).lower()
    
                  #if 'habits' in text:
                  for t in text.split('\r'):
                      if 'habits' in t:
                      
                        data_habits.append([j.split('.')[0],t.replace('\n',"").split(':')[1]])
                        
                  #if 'social history' in text:
                  for t in text.split('\r'):
                      if 'social history' in t:
                    
                        data_habits.append([j.split('.')[0],t.replace('\n',"").split(':')[1]])
             except:
                for p in doc.first_section.body.tables[2]:
                    row = p.as_row()
                    text = row.first_cell.to_string(aw.SaveFormat.TEXT).lower()

                    #if 'habits' in text:
                    for t in text.split('\r'):
                      if 'habits' in t:
                          data_habits.append([j.split('.')[0],t.replace('\n',"").split(':')[1]])
                        
                    #if 'social history' in text:
                    for t in text.split('\r'):
                      if 'social history' in t:
                          data_habits.append([j.split('.')[0],t.replace('\n',"").split(':')[1]])
        else:
            
            doc = textract.process(PATH_DOCS+j)
            comment = str(doc).lower().replace('|',"").split('\\n')
            for t in comment:
                if 'habits:' in t:
                    data_habits.append([j.split('.')[0],t.replace('\n',"").split(':')[1]])
                if 'social history:' in t:
                    data_habits.append([j.split('.')[0],t.replace('\n',"").split(':')[1]])\

    return data_habits
    
#FUNCTION TO GET THE KPS (KARNOSFKY PERFORMANCE) OF THE PATIENT 
def get_data_kps(dir_list,PATH_DOCS):

  data_kps = []

  for j in dir_list:
    if j.split('.')[1]!='pdf':
        doc = aw.Document(PATH_DOCS+j)
        formFields = doc.range.form_fields
    
        try:
            for p in doc.first_section.body.tables[1]:
                row = p.as_row()
   
                text = row.first_cell.to_string(aw.SaveFormat.TEXT).lower()
        
                for t in text.split('\r'):
                    if "kps" in t:
                        if 'physical examination:' in t:
                            data_kps.append([j.split('.')[0],t.replace('\n',"").replace('\u2028','.').replace('physical examination:','')])
                    
                    if 'karnofsky' in t:
                        if 'physical examination:' in t:
                            data_kps.append([j.split('.')[0],t.replace('\n',"").replace('\u2028','.').replace('physical examination:','')])
        except:
            print('error in '+j)
    else:     
        doc = textract.process(PATH_DOCS+j)
        comment = str(doc).lower().replace('|',"").split('\\n')
        for t in comment:
            if "kps" in t:
                if 'physical examination:' in t:
                    data_kps.append([j.split('.')[0],t.replace('\n',"").replace('\u2028','.').replace('physical examination:','')])
                    
            if 'karnofsky' in t:
                if 'physical examination:' in t:
                    data_kps.append([j.split('.')[0],t.replace('\n',"").replace('\u2028','.').replace('physical examination:','')])
 
  return data_kps
    
    
    
#GENERAL FUNCTION TO READ SOME KEYWORD FROM THE DOCUMENTS
def get_general_data_from_notes(dir_list,key_word,PATH_DOCS):

  data = []

  for j in dir_list:
    if j.split('.')[1]!='pdf':
        doc = aw.Document(PATH_DOCS+j)
        formFields = doc.range.form_fields
    
        try:
            for p in doc.first_section.body.tables[1]:
                row = p.as_row()
   
                text = row.first_cell.to_string(aw.SaveFormat.TEXT).lower()
        
                for t in text.split('\r'):
                    if keyword in t:
                        #if 'physical examination:' in t:
                        data.append([j.split('.')[0],t.replace('\n',"").replace('\u2028','.').replace(keyword,'')])                 
         except:
            print('ERROR IN PATIENT'+j+' CHECK DOCUMENT FORMAT OR KEYWORD')
    else:     
        doc = textract.process(PATH_DOCS+j)
        comment = str(doc).lower().replace('|',"").split('\\n')
        for t in comment:
            if keyword in t:
                
              data.append([j.split('.')[0],t.replace('\n',"").replace('\u2028','.').replace(keyword,'')])
   
  return data
  
  
###################


def get_weight_data_nut_docs(dir_list,PATH_DOCS_N):
  weight_data = []

  for j in dir_list:    

    doc2 = textract.process(PATH_DOCS_N+j)
    
    weight_info = []
    for p in str(doc2.lower()).split('\\n'):
        if 'poids / weight' in p:
            print(p)
            weight = p.split('poids / weight')[-1].replace('|',"")
            
        
        if 'poids habituel / usual weight' in p:
            #print(p)
            weight_usual = p.split('poids habituel / usual weight')[-1].replace('|',"")
        
        if 'taille / height' in p:
            height = p.split('taille / height')[-1].replace('|',"")
            #height_data.append([j.split('.')[0],height])
     
    weight_data.append([j.split('.')[0],height,weight,weight_usual])
  return weight_data
        



####################################################################################3

#FUNCTIONS TO READ THE REPLANNING REQUEST DOCUMENTS

noticed_by = {
  1:"RTT",
  2: "MD"}
noticed_during = {
  1:"Offline Review",
   2:"Daily setup"
}

noticed_by_i = {
  1:'[230,250];[60,75]',
  2:'[320,350];[60,75]'}

noticed_during_i = {
  1:'[80,100];[85,100]',
  2:'[220,280];[85,100]'}

replan_reason = {
  1:"Change in target or nodal volume",
  2: "Weight loss",
  3:"Change in skin separation",
   4:"Alteration in muscle mass and fat distribution",
    5:"Fluid shifts within body",
    6:"Change in patient immobilization device",
    7:"On-going difficulties in setup",
    8:"Other reason"
}

replan_reason_i = {
  1:[0,50],
  2:[51,100],
  3:[101,125],
    4:[126,150],
    5:[151,175],
    6:[176,200],
    7:[201,225],
    8:[226,250]
}

weight_lbs_or_kgs = {1:"lbs",2:"kg"}
weight_lbs_or_kgs_i = {1:[300,400],2:[401,500]}





def get_contours(img_bin):
    contours, _ = cv2.findContours(img_bin, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    return contours

def get_area_box(contours):
    contours2 = []
    contours3 = []
    for contour in contours:
        area = cv2.contourArea(contour)
        if (area)>=90000:
            contours2.append(contour)
    cnt = sorted(contours2, key=lambda x: cv2.contourArea(x),reverse=True)
    return cnt[0]

def get_mark(contours):
    contours2 = []
    for contour in contours:
        area = cv2.contourArea(contour)
        if 1000>area>=50:
        #contours3.append(contour)
            contours2.append(contour)
    return contours2

def get_middle(contours3,im55):
    edgs2 = []
    for j in contours3:
        xs = (np.array(j)[:,0][:,0])
        ys = (np.array(j)[:,0][:,1])
        edg1_y =np.max(ys)
        edg2_y = np.min(ys)
        edg1_x =np.max(xs)
        edg2_x = np.min(xs)
        
        edgs = []
    #edgs.append()
        edgs.append([(edg1_x+edg2_x)//2,(edg1_y+edg2_y)//2])
        for j in range(1,3):
            edgs.append([(edg1_x+edg2_x)//2+j,(edg1_y+edg2_y)//2-j])
            edgs.append([(edg1_x+edg2_x)//2-j,(edg1_y+edg2_y)//2+j])
            edgs.append([(edg1_x+edg2_x)//2+j,(edg1_y+edg2_y)//2+j])
            edgs.append([(edg1_x+edg2_x)//2-j,(edg1_y+edg2_y)//2-j])
            edgs.append([(edg1_x+edg2_x)//2-j,(edg1_y+edg2_y)//2])
            edgs.append([(edg1_x+edg2_x)//2+j,(edg1_y+edg2_y)//2])
            edgs.append([(edg1_x+edg2_x)//2,(edg1_y+edg2_y)//2-j])
            edgs.append([(edg1_x+edg2_x)//2,(edg1_y+edg2_y)//2+j])
        edgs2.append(edgs)
            
    #print(edgs)
    ctrs4 = []
    edgs5 = []
    
    for p in edgs2:
        edg4 = []
        for j in p:
            if (im55[j[1]][j[0]])!=0:
                ctrs4.append(j)
                edg4.append(j)
        edgs5.append(edg4)
        
    return ctrs4,edgs5
 

def get_read_images_boxes_replan(img):
    min_x=0
    max_x = 0
    min_y=0
    max_y=0
    
    _, img_bin = cv2.threshold(cv2.cvtColor(img, cv2.COLOR_BGR2GRAY), 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
    img_bin = 255 - img_bin
    CANNY_KERNEL_SIZE = 100
    img_canny = cv2.Canny(img_bin, CANNY_KERNEL_SIZE, CANNY_KERNEL_SIZE)

    contours = get_contours(img_canny)
    contours2 =  get_area_box(contours)
    
    im2 = img.copy()
    min_x = np.min(contours2[:,0][:,0])
    max_x = np.max(contours2[:,0][:,0])

    max_y = np.max(contours2[:,0][:,1])
    min_y = np.min(contours2[:,0][:,1])

    
    im6 = img_canny.copy()[min_y+1:max_y,min_x:max_x]
    #im7 = im6.copy()[:,1:25]
    im4 = im2.copy()[min_y+1:max_y,min_x:max_x]
    #im5 = im4.copy()[:,1:25]
    
    im44 = img_bin.copy()[min_y+1:max_y,min_x:max_x]
    #im55 = im44.copy()[:,1:25]
    
    return im44
    
tagRe = re.compile(r'\\x.*?(2)')

def read_checkboxes_reason(dir_list,PATH_R_DOCS_2):
    
    data = []
    #os.mkdir('temp2')
    for j in dir_list:
        specification = ""
        #from PyPDF2 import PdfReader
        #generate_pdf(PATH_R_DOCS_2+j, 'temp2') 
  
        word_reader = textract.process(PATH_R_DOCS_2+j)
      
        #pdf_reader.get_fields()
  
        doc = aw.Document(PATH_R_DOCS_2+j)

        extractedPage = doc.extract_pages(0, 1)
        extractedPage.save(f"Output_{1}_"+j.split('.')[0]+'.jpg')
          
        img=mpimg.imread('Output_1_'+j.split('.')[0]+'.jpg')
    
        image_to_read = get_read_images_boxes_replan(img)
      
        im2 = img.copy()
        im44 = get_read_images_boxes_replan(im2)
    
        contours5 = get_contours(im44)
    
        mark_ctrs = get_mark(contours5)
    
        ctrs4,edgs2 = get_middle(mark_ctrs,im44)

        ctrs4.sort(key=lambda ctrs4: ctrs4[1])

    
        limits_xy_reason = []
        limits_xy_kg_lbs = []
      
        for m in edgs2:
            if len(m)!=0:
                y1 =np.max(np.array(m)[:,1])
                y0 =np.min(np.array(m)[:,1])
                x1 =np.max(np.array(m)[:,0])
                x0  =np.min(np.array(m)[:,0])
                if 0<=x1<=100:
                    limits_xy_reason.append([[x0,x1],[y0,y1]])
                else:
                    limits_xy_kg_lbs.append([[x0,x1],[y0,y1]])
        reasons0 =[]
        for lim in limits_xy_reason:
            for ii in range(1,len(replan_reason_i)+1):
                if replan_reason_i[ii][0]<=lim[1][0]<=replan_reason_i[ii][1]:
                
                    reasons = replan_reason.get(ii)
                    page = word_reader
                
                    extract_text = str(page).split('REASON FOR REQUEST:')[-1].split('PROFILE:')[0]
                              
                    specification = extract_text.split(replan_reason.get(ii))[-1].split(str(replan_reason.get(ii+1)))[0]
               
                    specification = tagRe.sub('', specification)
                
                    specification = specification.replace('Specify:','')
                
                    specification = specification.replace('\\n','')
                
                    specification = specification.replace('Indicate source of new information (e.x. new imaging or lab tests):',"")
                    specification = specification.strip()
                    if specification=='.':
                        specification = specification.replace('.','')
                  
                    if reasons==replan_reason.get(2):
                        if len(limits_xy_kg_lbs)!=0:
                            for limm in limits_xy_kg_lbs:
                                for pp in range(1,len(weight_lbs_or_kgs_i)+1):
                                    if weight_lbs_or_kgs_i[pp][0]<=limm[0][0]<=weight_lbs_or_kgs_i[pp][1]:
                                        unit = weight_lbs_or_kgs.get(pp)
                                        specification_new = specification.replace('lbs or  kg',unit)
                                        reasons0.append([reasons,specification_new])
                        else:
                            specification_new = specification.replace('lbs or  kg','')
                            specification_new = specification_new.strip()
                            reasons0.append([reasons,specification_new])
                                
                    else:    
                        reasons0.append([reasons,specification])
 
      data.append([j.split('.')[0],reasons0]
      
      os.remove(f'Output_{0 + 1}_'+j.split('.')[0]+'.jpg')
      
    return  data
    
  
def get_area_box_v2(contours):
    contours2 = []
    contours3 = []
    for contour in contours:
        area = cv2.contourArea(contour)
        if 90000>=(area)>=10000:
            contours2.append(contour)
    cnt = sorted(contours2, key=lambda x: cv2.contourArea(x),reverse=True)
    return cnt[0]
      
def read_checkboxes_who_when_noticed(dir_list,PATH_R_DOCS_2):
  data = []
  for j in dir_list:
      specification = ""
      generate_pdf(PATH_R_DOCS_2+j, 'temp2') 
  
      #pdf_reader = PdfReader(open('temp2/'+j.split('.')[0]+'.pdf', "rb"))
      word_reader = textract.process(PATH_R_DOCS_2+j)
      #pdf_reader.get_fields()
    
      min_x=0
      max_x = 0
      min_y=0
      max_y=0
    #os.remove("/data2/odette/aria_Replan_request/temp/"+j.split('.')[0]+'.pdf')
      doc = aw.Document(PATH_R_DOCS_2+j)

    
      extractedPage = doc.extract_pages(0, 1)
      extractedPage.save(f"Output_{1}_"+j.split('.')[0]+'.jpg')
      img=mpimg.imread('Output_1_'+j.split('.')[0]+'.jpg')
    
      _, img_bin = cv2.threshold(cv2.cvtColor(img, cv2.COLOR_BGR2GRAY), 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
      img_bin = 255 - img_bin
      CANNY_KERNEL_SIZE = 100
      img_canny = cv2.Canny(img_bin, CANNY_KERNEL_SIZE, CANNY_KERNEL_SIZE)

      contours = get_contours(img_canny)
      contours2 =  get_area_box_v2(contours)
    
      im2 = img.copy()
      min_x = np.min(contours2[:,0][:,0])
      max_x = np.max(contours2[:,0][:,0])

      max_y = np.max(contours2[:,0][:,1])
      min_y = np.min(contours2[:,0][:,1])

    #im10 = cv2.drawContours(im2.copy(), contours2, -1, (0,255,0), 3)
      im6 = img_canny.copy()[min_y+1:max_y,min_x:max_x]
      #im7 = im6.copy()[0:80,:] #upper part
      im4 = im2.copy()[min_y+1:max_y,min_x:max_x]
      #im5 = im4.copy()[0:80,:] # upper part

    
      im44 = img_bin.copy()[min_y+1:max_y,min_x:max_x]
    
      contours3 = get_contours(im44)
      mark_ctrs = get_mark(contours3)
    
      ctrs4,edgs2 = get_middle(mark_ctrs,im44)
      ctrs4.sort(key=lambda ctrs4: ctrs4[1])
    
      im55 = im4.copy()
      im10 = cv2.drawContours(im55.copy(), mark_ctrs, -1, (0,255,0), 3)
    

      limits_xy = []
      for m in edgs2:
        if len(m)!=0:
            y1 =np.max(np.array(m)[:,1])
            y0 =np.min(np.array(m)[:,1])
            x1 = np.max(np.array(m)[:,0])
            x0 = np.min(np.array(m)[:,0])
            limits_xy.append([[x0,x1],[y0,y1]])
    
      who_noticed =[]
    
#    print(limits_xy)
      for lim in limits_xy:
          for ii in range(1,len(noticed_by_i)+1):
              coords_x = noticed_by_i.get(ii).split(';')[0].replace('[','').replace(']','').split(',')
              coords_y = noticed_by_i.get(ii).split(';')[1].replace('[','').replace(']','').split(',')
              if int(coords_x[0])<=lim[0][0]<=int(coords_x[1]):
                
                  if int(coords_y[0])<=lim[1][0]<=int(coords_y[1]):
                      noticed = noticed_by.get(ii)
                      who_noticed.append(noticed)
    
    
      page = word_reader
      extract_text = str(page).replace('\n',"").split('Anatomical changes first noticed by:')[-1].split('during:')[0]
      specification = tagRe.sub('', extract_text)
      specification = specification.split('(specify)')[-1]
      #specification = specification.replace('specify','')
      specification = specification.replace('\\n','')
      specification = specification.strip()
      who_noticed.append(specification)
    
      during = []
      for lim in limits_xy:
          for ii in range(1,len(noticed_during_i)+1):
              coords_x = noticed_during_i.get(ii).split(';')[0].replace('[','').replace(']','').split(',')
              coords_y = noticed_during_i.get(ii).split(';')[1].replace('[','').replace(']','').split(',')
              if int(coords_x[0])<=lim[0][1]<=int(coords_x[1]):
                  if int(coords_y[0])<=lim[1][0]<=int(coords_y[1]):
                      during_text = noticed_during.get(ii)
                      during.append(during_text)  
                      
      extract_text_2 = str(page).replace('\n',"").split('Anatomical changes first noticed by:')[-1].split('during:')[-1].split('COMMENTS / INSTRUCTIONS:')[0]
      specification_2 = tagRe.sub('', extract_text_2)
      specification_2 = specification_2.split('(specify)')[-1]
      #specification_2 = extract_text_2.split('(specify)')[-1]
      specification_2 = specification_2.replace('\\n','')
      specification_2 = specification_2.strip()
      #specification_2 = specification_2.replace('\u2002\u2002\u2002\u2002\u2002',"")
            
      during.append(specification_2)
      
    os.remove(f'Output_{1}_'+j.split('.')[0]+'.jpg')
    
    data.append([j.split('.')[0],during,who_noticed]
    
    return data


def read_fractions_remaining_and_total_dose(dir_list,PATH_R_DOCS):
    fx_request = []
    for j in dir_list:
   
        doc = aw.Document(PATH_R_DOCS+j)
        formFields = doc.range.form_fields
            #try:
   
        total_dose = []
        fx_value = []
        for p in doc.first_section.body.tables[0]:
            row = p.as_row()
            text = row.first_cell.to_string(aw.SaveFormat.TEXT).lower()
        #print(text)
        
        
            for t in text.split('\r'):
                if 'total dose:' in t:
                    if 'gy' in t.split(':')[1]:
                    #total_dose.append([j.split('.')[0],t.split(':')[1].replace('gy','')])
                      total_dose.append(t.split(':')[1].replace('gy','').replace(' ',""))
                    else:
                    #total_dose.append([j.split('.')[0],t.split(':')[1]])
                      total_dose.append(t.split(':')[1].replace(' ',""))     
        
            for t in text.split('\r'):
                if 'number of fractions remaining:' in t:
                    if '/' in t.split(':')[1]:
                    #fx_value.append([j.split('.')[0],t.split(':')[1].split('/')[1].replace('fx',"").replace(" ","")])
                        fx_value.append(t.split(':')[1].split('/')[1].replace('fx',"").replace(" ",""))
                    else:
                    #fx_value.append([j.split('.')[0],t.split(':')[1].replace('fx',"").replace(" ","")])
                        fx_value.append(t.split(':')[1].replace('fx',"").replace(" ",""))
                    
      for l in range(0,len(total_dose)):
          fx_request.append([j.split('.')[0],total_dose[l],fx_value[l]])
          
   return fx_request

#FUNCTION TO READ WHO SIGNED THE DOCUMENT, I.E. TO GET THE NAME OF THE PHYSICIT OR MD OR RTT THAT SIGNED
def get_signature(dir_list,PATH_R_DOCS_2):
    signature_data = []

    for j in dir_list:
    
    doc = textract.process(PATH_R_DOCS_2+j)
    comment = str(doc).lower().replace('|',"").split('\\n')
   
    for p in comment:
        if "electronically signed by" in p:
            name = p.split('electronically signed by')[-1].replace("'","").replace('"',"")
            signature_data.append([j.split('.')[0],name])
            
    return signature_data
                

#FUNCTION TO GET THE INSTRUCTIONS/COMMENTS SECTION 
def read_comments_instructions_replan_docs(dir_list,PATH_R_DOCS):
  comments = []
  for j in dir_list:
      doc = textract.process(PATH_R_DOCS+j)
   
      comment = ('.').join(str(doc).split('COMMENTS / INSTRUCTIONS')[1].split('\\n')[1:3]).replace('|',"")
 
      comment_2 = (('.').join(str(doc).split('COMMENTS / INSTRUCTIONS')[1].split('\\n')[1:3]).replace('|',""))
      comment_3 = comment_2.split('.')
 
      if comment_3[0]=='':
          comments.append(comment_2[1:])
      else:
          comments.append(comment_2)
  return comments
 
#FUNCTION TO GET THE DATE FROM THE DOCUMENTS. 
#NOTE THAT IF THE DOCUMENTS ARE IN .DOCX. THESE ARE READ BY TEXTRACT
#IF THE DOCUMENT IS .DOC, THE DOCUMENT IS CONVERTED AND SAVED IN .DOCX TO BE READ BY TEXTRACT
#THE FUNCTION ASSUMES THAT THE DOCUMENTS ARE IN WORD DOCUMENT FORMAT
def replan_date(dir_list,PATH_R_DOCS):
  if os.path.isdir('temp')==False:
        os.mkdir('temp')
  #os.mkdir('temp')
  data_request_replan_date = []


  for j in dir_list:
    if j.split('.')[-1]=='docx':
        doc = textract.process(PATH_R_DOCS+j)
        comment = str(doc).lower().replace('|',"").split('\\n')
      
        for p in comment:
            if 'date:' in p:
                #print(p)
                data_request_replan_date.append([j.split('.')[0],p.split(':')[-1].replace('|','')])
       
    elif j.split('.')[-1]=='doc':
    
        generate_docx(PATH_R_DOCS+j, 'temp')
        
        doc = textract.process("temp/"+j.split('.')[0]+'.docx')
        comment = str(doc).lower().replace('|',"").split('\\n')
        for p in comment:
            if 'date:' in p:
                #print(p)
                data_request_replan_date.append([j.split('.')[0],p.split(':')[-1].replace('|','')])

        os.remove("temp/"+j.split('.')[0]+'.docx') #REMOVE DOCUMENT FROM THE TEMPORARY FOLDER
        
    return data_request_replan_date

#THIS FUNCTION GETS THE SPECIFICATION OF REQUEST REPLAN DOCUMENT. i.e., THE REMAINING FRACTIONS
#TO READ THIS TEXTBOX IN THE DOCUMENT IS BETTER TO CONVERT THEM TO .PDF
#THE FUNCTION CREATES A TEMPORAL FOLDER TO KEEP ALL THE DOCUMENTS IN .PDF
#OTHER APPROACH WOULD BE READ ALL THE DOCUMENTS IN .PDF, THUS CONVERT THEM ALL OF THEM, HOWEVER DOES NOT GURANTEE 
#APPROPIATE READING FOR SOME INFORMATION
def get_specifications(PATH_R_DOCS_2):
  if os.path.isdir('temp2')==False:
        os.mkdir('temp2')
  #os.mkdir('temp2' #CREATES THE FOLDER
  
  data_specifications_dose = []
  number_specification_fractions = []
  dir_list = os.listdir(PATH_R_DOCS_2)

  for j in dir_list:    
      generate_pdf(PATH_R_DOCS_2+j, 'temp2') 
  
      pdf_reader = PdfReader(open('temp2/'+j.split('.')[0]+'.pdf', "rb"))

      page = pdf_reader.pages[0]
      
      for p in str(page.extract_text()).lower().split('\n'):
          if 'total dose:' in p:
              if 'specify:' in p:
                  text = p.replace('\u2002\u2002\u2002\u2002\u2002 ','')#ERASES SOME EXTRA INFO/SYMBOLS NOT WANTED
                  text2 = text.split('specify:')[-1]
                  
                  data_specifications_dose.append([j.split('.')[0],text2])
                  
          if 'number of fractions remaining:' in p:
              if 'specify:' in p:
                  
                  text = p.replace('\u2002\u2002\u2002\u2002\u2002 ','')
                  text2 = text.split('specify:')[-1]
                  
                  number_specification_fractions.append([j.split('.')[0],text2]) 
                
      os.remove("temp2/"+j.split('.')[0]+'.pdf')
      
  return data_specifications_dose,number_specification_fractions
  
#SAVE THE DATA IN .CSV FILE 
#TO DO: SAVE IT IN .JSON
def save_docsv2(ids_mrn,data):
  data_w_id = []
  for j in ids_mrn:
      for k in data:
          if str(j[1]) in k[0]:
              data_w_id.append([j[0],k[0],k[1],k[2]])
  comment_data = pd.DataFrame(data_w_id ,columns=['id','mrn','comment'])
  comment_data.to_csv('OUTPUT_DOC.csv',index=False)
  
  return
  
def save_docs(ids_mrn,data):
  data_w_id = []
  for j in ids_mrn:
      for k in data:
          if str(j[1]) in k[0]:
              data_w_id.append([j[0],k[0],k[1]])
  comment_data = pd.DataFrame(data_w_id ,columns=['id','mrn','comment'])
  comment_data.to_csv('OUTPUT_DOC.csv',index=False)
  
  return 

###CHANGE HERE THE PATHS 
PATH_DOCS = '/data2/odette/aria_RADO_DOCS/RADO_ONCS/'
PATH_R_DOCS = '/mnt/iDriveShare/Yossera_HNC_Summer/Replan_new_documents/'
PATH_DOCS_N = '/data2/odette/aria_RADO_DOCS/NUT_ONCS/'

dir_list = os.listdir(PATH_DOCS)
dir_list_R = os.listdir(PATH_R_DOCS)
dir_list_N = os.listdir(PATH_DOCS_N)
dir_list_R = os.listdir(PATH_R_DOCS)

data_ids_mrn = pd.read_csv('/mnt/iDriveShare/Kayla/for_yossera_new_patient_list.csv')

#GETS THE MRN FOR EACH PATIENT ID
ids_mrn = []
for j in data_ids_mrn[['id','mrn']].values:
    if list(j) not in ids_mrn:
      ids_mrn.append(list(j))

#CREATE DICTIONARY

df = pd.DataFrame(np.array(ids_mrn)[:,1],index = np.array(ids_mrn)[:,0],columns=['mrn'])
dict_patients = df.to_dict('index')


