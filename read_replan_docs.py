import textract
from docx import Document
import os
import subprocess
import aspose.words as aw
from json import loads, dumps
import matplotlib.image as mpimg
import re
import cv2 
import numpy as np
import pandas as pd
import textract
import PyPDF2
import json
from pdf2image import convert_from_path
from PyPDF2 import PdfReader



def generate_docx(doc_path, path):

    subprocess.call(['soffice',
                 # '--headless',
                 '--convert-to',
                 'docx',
                 '--outdir',
                 path,
                 doc_path])
    return doc_path


def generate_pdf(doc_path, path):

    subprocess.call(['soffice',
                 # '--headless',
                 '--convert-to',
                 'pdf',
                 '--outdir',
                 path,
                 doc_path])
    return doc_path


noticed_by = {
  1:"RTT",
  2: "MD"}
noticed_during = {
  1:"Offline Review",
   2:"Daily setup"
}

noticed_by_i = {
  1:'[230,250];[60,75]',
  2:'[320,350];[60,75]'}

noticed_during_i = {
  1:'[80,100];[85,100]',
  2:'[220,280];[85,100]'}

replan_reason = {
  1:"Change in target or nodal volume",
  2: "Weight loss",
  3:"Change in skin separation",
   4:"Alteration in muscle mass and fat distribution",
    5:"Fluid shifts within body",
    6:"Change in patient immobilization device",
    7:"On-going difficulties in setup",
    8:"Other reason"
}

replan_reason_i = {
  1:[0,50],
  2:[51,100],
  3:[101,125],
    4:[126,150],
    5:[151,175],
    6:[176,200],
    7:[201,225],
    8:[226,250]
}

weight_lbs_or_kgs = {1:"lbs",2:"kg"}
weight_lbs_or_kgs_i = {1:[300,400],2:[401,500]}



#plt.figure()
import aspose.words as aw
import matplotlib.image as mpimg


def get_images_range(img_bin,r=0.5):
    im44 = img_bin.copy()[int(size_y*r):,600:1400][:,200:250]
    #plt.imshow(img_bin[:,600:1400])
    #plt.show()
    contours = get_contours(im44)
    contours2 =  get_area_box_v3(contours)
    #if 
    #print(contours2)
    edgs0 = []
    for p in contours2:
        xs = (np.array(p)[:,0][:,0])
        ys = (np.array(p)[:,0][:,1])
        edg1_y =np.max(ys)
        edg2_y = np.min(ys)
        edg1_x =np.max(xs)
        edg2_x = np.min(xs)    
        edgs0.append(edg1_y)
        edgs0.append(edg2_y)
    
    #if len(edgs0)!=0
    if len(edgs0)==0:
        edgs0,r = get_images_range(img_bin,r*0.5)
    else:
        return edgs0,r
    return edgs0,r

recommendations = {
  1:"Schedule a CT scan and replan",
  2: "Continue with current plan"}

recommendations_i = {
  1:[0,30],
  2:[35,60]}
  
replan_or_not = {1: "REPLAN", 2:"NO REPLAN"}


def get_area_box_v3(contours):
    contours2 = []
    contours3 = []
    for contour in contours:
        area = cv2.contourArea(contour)
        #print(area)
        if 700>=(area)>=500:
            contours2.append(contour)
    cnt = sorted(contours2, key=lambda x: cv2.contourArea(x),reverse=True)
    return cnt



def get_contours(img_bin):
    contours, _ = cv2.findContours(img_bin, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    return contours

def get_area_box(contours):
    contours2 = []
    contours3 = []
    for contour in contours:
        area = cv2.contourArea(contour)
        if (area)>=90000:
            contours2.append(contour)
    cnt = sorted(contours2, key=lambda x: cv2.contourArea(x),reverse=True)
    return cnt[0]

def get_mark(contours):
    contours2 = []
    for contour in contours:
        area = cv2.contourArea(contour)
        if 1000>area>=50:
        #contours3.append(contour)
            contours2.append(contour)
    return contours2

def get_middle(contours3,im55):
    edgs2 = []
    for j in contours3:
        xs = (np.array(j)[:,0][:,0])
        ys = (np.array(j)[:,0][:,1])
        edg1_y =np.max(ys)
        edg2_y = np.min(ys)
        edg1_x =np.max(xs)
        edg2_x = np.min(xs)
        
        edgs = []
    #edgs.append()
        edgs.append([(edg1_x+edg2_x)//2,(edg1_y+edg2_y)//2])
        for j in range(1,3):
            edgs.append([(edg1_x+edg2_x)//2+j,(edg1_y+edg2_y)//2-j])
            edgs.append([(edg1_x+edg2_x)//2-j,(edg1_y+edg2_y)//2+j])
            edgs.append([(edg1_x+edg2_x)//2+j,(edg1_y+edg2_y)//2+j])
            edgs.append([(edg1_x+edg2_x)//2-j,(edg1_y+edg2_y)//2-j])
            edgs.append([(edg1_x+edg2_x)//2-j,(edg1_y+edg2_y)//2])
            edgs.append([(edg1_x+edg2_x)//2+j,(edg1_y+edg2_y)//2])
            edgs.append([(edg1_x+edg2_x)//2,(edg1_y+edg2_y)//2-j])
            edgs.append([(edg1_x+edg2_x)//2,(edg1_y+edg2_y)//2+j])
        edgs2.append(edgs)
            
    #print(edgs)
    ctrs4 = []
    edgs5 = []
    
    for p in edgs2:
        edg4 = []
        for j in p:
            if (im55[j[1]][j[0]])!=0:
                ctrs4.append(j)
                edg4.append(j)
        edgs5.append(edg4)
        
    return ctrs4,edgs5
 

def get_read_images_boxes_replan(img):
    min_x=0
    max_x = 0
    min_y=0
    max_y=0
    
    _, img_bin = cv2.threshold(cv2.cvtColor(img, cv2.COLOR_BGR2GRAY), 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
    img_bin = 255 - img_bin
    CANNY_KERNEL_SIZE = 100
    img_canny = cv2.Canny(img_bin, CANNY_KERNEL_SIZE, CANNY_KERNEL_SIZE)

    contours = get_contours(img_canny)
    contours2 =  get_area_box(contours)
    
    im2 = img.copy()
    min_x = np.min(contours2[:,0][:,0])
    max_x = np.max(contours2[:,0][:,0])

    max_y = np.max(contours2[:,0][:,1])
    min_y = np.min(contours2[:,0][:,1])

    
    im6 = img_canny.copy()[min_y+1:max_y,min_x:max_x]
    #im7 = im6.copy()[:,1:25]
    im4 = im2.copy()[min_y+1:max_y,min_x:max_x]
    #im5 = im4.copy()[:,1:25]
    
    im44 = img_bin.copy()[min_y+1:max_y,min_x:max_x]
    #im55 = im44.copy()[:,1:25]
    
    return im44
    
tagRe = re.compile(r'\\x.*?(2)')


def read_checkboxes_recommendation(PATH_R_DOCS_2,j):
 
  generate_pdf(PATH_R_DOCS_2+j, 'temp2') 
  
  reader = PdfReader(open('temp2/'+j.split('.')[0]+'.pdf', "rb"))

  page = 0
  
  images = convert_from_path('temp2/'+j.split('.')[0]+'.pdf')
    
  String = "Recommendations"

  # Extract text and do the search
  for i in range(0, len(reader.pages)):
      PageObj = reader.pages[i]
      Text = PageObj.extract_text()
      if re.search(String,Text):
          page = i       
    
  images[page].save(f'output_'+j.split('.')[0]+'.jpg', 'JPEG')
  img2=mpimg.imread(f'output_'+j.split('.')[0]+'.jpg')
        
  size_y = len(img2)
  _, img_bin = cv2.threshold(cv2.cvtColor(img2, cv2.COLOR_BGR2GRAY), 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
  img_bin = 255 - img_bin
  CANNY_KERNEL_SIZE = 100
  img_canny = cv2.Canny(img_bin, CANNY_KERNEL_SIZE, CANNY_KERNEL_SIZE)

  edgs0,r = get_images_range(img_bin)

  y0 = np.min(edgs0)
  y1 = np.max(edgs0)

  im6 = img2.copy()[int(size_y*r):,600:1400][y0-2:y1+1,200:250]
  im66 = img_bin.copy()[int(size_y*r):,600:1400][y0-2:y1+1,200:250]
  contours4 = get_contours(im66)
  contours5 =  get_area_box_v3(contours4)
  mark_ctrs = get_mark(contours5)
  ctrs4,edgs2 = get_middle(mark_ctrs,im66)
    
  ctrs4.sort(key=lambda ctrs4: ctrs4[1])
   
    
  limits_xy = []
  for m in edgs2:
      if len(m)!=0:
          y1 =np.max(np.array(m)[:,1])
          y0 =np.min(np.array(m)[:,1])
          x1 = np.max(np.array(m)[:,0])
          x0 = np.min(np.array(m)[:,0])
          limits_xy.append([[x0,x1],[y0,y1]])
            
  recomm = ""
  R_NR = ""
  recommendation = []
  for lim in limits_xy:
      for ii in range(1,len(recommendations_i)+1):
          coords_y = recommendations_i.get(ii)  
          if int(coords_y[0])<=lim[1][0]<=int(coords_y[1]):  
              recomm = recommendations.get(ii)
              R_NR = replan_or_not.get(ii)
                #recommendation.append(recomm,R_NR)
                
  os.remove(f'output_'+j.split('.')[0]+'.jpg')
    
  return recomm,R_NR

def read_checkboxes_reason(PATH_R_DOCS_2,j):
    #dir_list = os.listdir(PATH_R_DOCS_2)
    #data = []
    #os.mkdir('temp2')
    #for j in dir_list:
    specification = ""
        #from PyPDF2 import PdfReader
        #generate_pdf(PATH_R_DOCS_2+j, 'temp2') 
  
    word_reader = textract.process(PATH_R_DOCS_2+j)
      
        #pdf_reader.get_fields()
  
    doc = aw.Document(PATH_R_DOCS_2+j)

    extractedPage = doc.extract_pages(0, 1)
    extractedPage.save(f"Output_{0 + 1}_"+j.split('.')[0]+'.jpg')
          
    img=mpimg.imread('Output_1_'+j.split('.')[0]+'.jpg')
    
    image_to_read = get_read_images_boxes_replan(img)
      
    im2 = img.copy()
    im44 = get_read_images_boxes_replan(im2)
    
    contours5 = get_contours(im44)
    
    mark_ctrs = get_mark(contours5)
    
    ctrs4,edgs2 = get_middle(mark_ctrs,im44)

    ctrs4.sort(key=lambda ctrs4: ctrs4[1])

    
    limits_xy_reason = []
    limits_xy_kg_lbs = []
      
    for m in edgs2:
        if len(m)!=0:
            y1 =np.max(np.array(m)[:,1])
            y0 =np.min(np.array(m)[:,1])
            x1 =np.max(np.array(m)[:,0])
            x0  =np.min(np.array(m)[:,0])
            if 0<=x1<=100:
                limits_xy_reason.append([[x0,x1],[y0,y1]])
            else:
                limits_xy_kg_lbs.append([[x0,x1],[y0,y1]])
    reasons0 =[]
    specifications0 = []
    for lim in limits_xy_reason:
        for ii in range(1,len(replan_reason_i)+1):
            if replan_reason_i[ii][0]<=lim[1][0]<=replan_reason_i[ii][1]:
                
                reasons = replan_reason.get(ii)
                page = word_reader
                
                extract_text = str(page).split('REASON FOR REQUEST:')[-1].split('PROFILE:')[0]
                              
                specification = extract_text.split(replan_reason.get(ii))[-1].split(str(replan_reason.get(ii+1)))[0]
                    #print(specification)
                specification = specification.replace('Specify:','')
                specification = tagRe.sub('', specification)
                specification = specification.replace('\\n','')
                
                specification = specification.replace('Indicate source of new information (e.x. new imaging or lab tests):',"")
                specification = specification.strip()
                    
                if specification=='.':
                    specification = specification.replace('.','')
                  
                if reasons==replan_reason.get(2):
                    if len(limits_xy_kg_lbs)!=0:
                        for limm in limits_xy_kg_lbs:
                            for pp in range(1,len(weight_lbs_or_kgs_i)+1):
                                if weight_lbs_or_kgs_i[pp][0]<=limm[0][0]<=weight_lbs_or_kgs_i[pp][1]:
                                    unit = weight_lbs_or_kgs.get(pp)
                                    specification_new = specification.replace('lbs or  kg',unit)
                                    reasons0.append(reasons)
                                    specifications0.append(specification_new )
                    else:
                        specification_new = specification.replace('lbs or  kg','')
                        specification_new = specification_new.strip()
                        reasons0.append(reasons)
                        specifications0.append(specification_new)
                                
                else:    
                    reasons0.append(reasons)
                    specifications0.append(specification)
 
    #data.append([j.split('.')[0],reasons0])
    os.remove(f'Output_{0 + 1}_'+j.split('.')[0]+'.jpg')
      
    return  reasons0,specifications0
    
def get_area_box_v2(contours):
    contours2 = []
    contours3 = []
    for contour in contours:
        area = cv2.contourArea(contour)
        if 90000>=(area)>=10000:
            contours2.append(contour)
    cnt = sorted(contours2, key=lambda x: cv2.contourArea(x),reverse=True)
    return cnt[0]
      
def read_checkboxes_who_when_noticed(PATH_R_DOCS_2,j):
    data = []
    #dir_list = os.listdir(PATH_R_DOCS_2)
 
    #for j in dir_list:
    specification = ""
        #generate_pdf(PATH_R_DOCS_2+j, 'temp2') 
  
      #pdf_reader = PdfReader(open('temp2/'+j.split('.')[0]+'.pdf', "rb"))
    word_reader = textract.process(PATH_R_DOCS_2+j)
      #pdf_reader.get_fields()
    
    min_x=0
    max_x = 0
    min_y=0
    max_y=0
    #os.remove("/data2/odette/aria_Replan_request/temp/"+j.split('.')[0]+'.pdf')
    doc = aw.Document(PATH_R_DOCS_2+j)
    extractedPage = doc.extract_pages(0, 1)
    extractedPage.save(f"Output_{1}_"+j.split('.')[0]+'.jpg')
    img=mpimg.imread('Output_1_'+j.split('.')[0]+'.jpg')
    
    _, img_bin = cv2.threshold(cv2.cvtColor(img, cv2.COLOR_BGR2GRAY), 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
    img_bin = 255 - img_bin
    CANNY_KERNEL_SIZE = 100
    img_canny = cv2.Canny(img_bin, CANNY_KERNEL_SIZE, CANNY_KERNEL_SIZE)

    contours = get_contours(img_canny)
    contours2 =  get_area_box_v2(contours)
    
    im2 = img.copy()
    min_x = np.min(contours2[:,0][:,0])
    max_x = np.max(contours2[:,0][:,0])

    max_y = np.max(contours2[:,0][:,1])
    min_y = np.min(contours2[:,0][:,1])

    im6 = img_canny.copy()[min_y+1:max_y,min_x:max_x]
        #im7 = im6.copy()[0:80,:] #upper part
    im4 = im2.copy()[min_y+1:max_y,min_x:max_x]
      #im5 = im4.copy()[0:80,:] # upper part
 
    im44 = img_bin.copy()[min_y+1:max_y,min_x:max_x]
    
    contours3 = get_contours(im44)
    mark_ctrs = get_mark(contours3)
    
    ctrs4,edgs2 = get_middle(mark_ctrs,im44)
    ctrs4.sort(key=lambda ctrs4: ctrs4[1])
    
    im55 = im4.copy()
    #im10 = cv2.drawContours(im55.copy(), mark_ctrs, -1, (0,255,0), 3)
    
    limits_xy = []
    for m in edgs2:
        if len(m)!=0:
            y1 =np.max(np.array(m)[:,1])
            y0 =np.min(np.array(m)[:,1])
            x1 = np.max(np.array(m)[:,0])
            x0 = np.min(np.array(m)[:,0])
            limits_xy.append([[x0,x1],[y0,y1]])
    
    who_noticed =[]
    
#    print(limits_xy)
    for lim in limits_xy:
        for ii in range(1,len(noticed_by_i)+1):
            coords_x = noticed_by_i.get(ii).split(';')[0].replace('[','').replace(']','').split(',')
            coords_y = noticed_by_i.get(ii).split(';')[1].replace('[','').replace(']','').split(',')
            if int(coords_x[0])<=lim[0][0]<=int(coords_x[1]):                
                if int(coords_y[0])<=lim[1][0]<=int(coords_y[1]):
                    noticed = noticed_by.get(ii)
                    who_noticed.append(noticed)

    page = word_reader
    extract_text = str(page).replace('\n',"").split('Anatomical changes first noticed by:')[-1].split('during:')[0]
    specification = tagRe.sub('', extract_text)
    specification = specification.split('(specify)')[-1]
      #specification = specification.replace('specify','')
    specification = specification.replace('\\n','')
    specification = specification.strip()
    #who_noticed.append(specification)
    
    during = []
    for lim in limits_xy:
        for ii in range(1,len(noticed_during_i)+1):
            coords_x = noticed_during_i.get(ii).split(';')[0].replace('[','').replace(']','').split(',')
            coords_y = noticed_during_i.get(ii).split(';')[1].replace('[','').replace(']','').split(',')
            if int(coords_x[0])<=lim[0][1]<=int(coords_x[1]):
                if int(coords_y[0])<=lim[1][0]<=int(coords_y[1]):
                    during_text = noticed_during.get(ii)
                    during.append(during_text)  
                      
    extract_text_2 = str(page).replace('\n',"").split('Anatomical changes first noticed by:')[-1].split('during:')[-1].split('COMMENTS / INSTRUCTIONS:')[0]
    specification_2 = tagRe.sub('', extract_text_2)
    specification_2 = specification_2.split('(specify)')[-1]
      #specification_2 = extract_text_2.split('(specify)')[-1]
    specification_2 = specification_2.replace('\\n','')
    specification_2 = specification_2.strip()
      #specification_2 = specification_2.replace('\u2002\u2002\u2002\u2002\u2002',"")
            
    #during.append(specification_2)
      
    os.remove(f'Output_{1}_'+j.split('.')[0]+'.jpg')
    #data.append([j.split('.')[0],during,who_noticed])
    
    return who_noticed,specification,during,specification_2


def read_fractions_remaining_and_total_dose(PATH_R_DOCS,j):
    fx_request = []
    #dir_list = os.listdir(PATH_R_DOCS)
    #for j in dir_list:
   
    doc = aw.Document(PATH_R_DOCS+j)
    formFields = doc.range.form_fields
            #try:
   
    total_dose = []
    fx_value = []
    for p in doc.first_section.body.tables[0]:
        row = p.as_row()
        text = row.first_cell.to_string(aw.SaveFormat.TEXT).lower()
        #print(text)
        for t in text.split('\r'):
            if 'total dose:' in t:
                if 'gy' in t.split(':')[1]:
                    #total_dose.append([j.split('.')[0],t.split(':')[1].replace('gy','')])
                    total_dose.append(t.split(':')[1].replace('gy','').replace(' ',""))
                else:
                #total_dose.append([j.split('.')[0],t.split(':')[1]])
                    total_dose.append(t.split(':')[1].replace(' ',""))     
        
        for t in text.split('\r'):
            if 'number of fractions remaining:' in t:
                if '/' in t.split(':')[1]:
                #fx_value.append([j.split('.')[0],t.split(':')[1].split('/')[1].replace('fx',"").replace(" ","")])
                    fx_value.append(t.split(':')[1].split('/')[1].replace('fx',"").replace(" ",""))
                else:
                #fx_value.append([j.split('.')[0],t.split(':')[1].replace('fx',"").replace(" ","")])
                    fx_value.append(t.split(':')[1].replace('fx',"").replace(" ",""))
                    
    #for l in range(0,len(total_dose)):
     #   fx_request.append([j.split('.')[0],total_dose[l],fx_value[l]])          
        
    return total_dose[0].rstrip(),fx_value[0].strip()

#FUNCTION TO READ WHO SIGNED THE DOCUMENT, I.E. TO GET THE NAME OF THE PHYSICIT OR MD OR RTT THAT SIGNED
def get_signature(PATH_R_DOCS_2,j):
    signature_data = []
    #dir_list = os.listdir(PATH_R_DOCS_2)

    #for j in dir_list:
    
    doc = textract.process(PATH_R_DOCS_2+j)
    comment = str(doc).lower().replace('|',"").split('\\n')
   
    for p in comment:
        if "electronically signed by" in p:
            name = p.split('electronically signed by')[-1].replace("'","").replace('"',"")
            signature_data.append([j.split('.')[0],name])
            
            return name.strip()
                

#FUNCTION TO GET THE INSTRUCTIONS/COMMENTS SECTION 
def read_comments_instructions_replan_docs(PATH_R_DOCS,j):
    comments = []
    #dir_list = os.listdir(PATH_R_DOCS)
    #for j in dir_list:
    doc = textract.process(PATH_R_DOCS+j)
   
    comment = str(doc).split('COMMENTS / INSTRUCTIONS:')[-1].split('Request discussed with Physics')[0]
    comment2 = comment.replace('\\n','').strip()
    #comment_2 = (('.').join(str(doc).split('COMMENTS / INSTRUCTIONS')[1].split('\\n')[1:3]).replace('|',""))
    #comment_3 = comment_2.split('.')
 
    #if comment_3[0]=='':
        #comment_4 = comment_2[1:]
    comment_3 = tagRe.sub('', comment2)
    #    comments.append(comment_4)
    #else:
     #   comment_4 = tagRe.sub('', comment_2)
      #  comments.append(comment_4)
        
    return comment_3.strip()

def get_comments_in_physics_section(PATH_DOCS_R,j):
    doc = textract.process(PATH_DOCS_R+j)
   
    comment = (str(doc).replace('\\n','').split('Request discussed with Physics')[-1]).split('Recommendations:')[0]
    
    comment_2 = comment.replace('\\n','')
    comment_3 = tagRe.sub('', comment_2[1:])
    #comment_3.strip()
    #comment_3 = comment_2.split('.')
    
    return comment_3.strip()
 
#FUNCTION TO GET THE DATE FROM THE DOCUMENTS. 
#NOTE THAT IF THE DOCUMENTS ARE IN .DOCX. THESE ARE READ BY TEXTRACT
#IF THE DOCUMENT IS .DOC, THE DOCUMENT IS CONVERTED AND SAVED IN .DOCX TO BE READ BY TEXTRACT
#THE FUNCTION ASSUMES THAT THE DOCUMENTS ARE IN WORD DOCUMENT FORMAT
def replan_date(PATH_R_DOCS,j):
    
    data_request_replan_date = []
    #dir_list = os.listdir(PATH_R_DOCS)
    #for j in dir_list:
    if j.split('.')[-1]=='docx':
        doc = textract.process(PATH_R_DOCS+j)
        comment = str(doc).lower().replace('|',"").split('\\n')
        for p in comment:
            if 'date:' in p:
                data_request_replan_date.append(p.split(':')[-1].replace('|','').strip())
       
    elif j.split('.')[-1]=='doc':    
        generate_docx(PATH_R_DOCS+j, 'temp')
        
        doc = textract.process("temp/"+j.split('.')[0]+'.docx')
        comment = str(doc).lower().replace('|',"").split('\\n')
        for p in comment:
            if 'date:' in p:
            #print(p)
                data_request_replan_date.append(p.split(':')[-1].replace('|','').strip())

        os.remove("temp/"+j.split('.')[0]+'.docx') #REMOVE DOCUMENT FROM THE TEMPORARY FOLDER
    #os.remove("temp")
    return data_request_replan_date[0]

#THIS FUNCTION GETS THE SPECIFICATION OF REQUEST REPLAN DOCUMENT. i.e., THE REMAINING FRACTIONS
#TO READ THIS TEXTBOX IN THE DOCUMENT IS BETTER TO CONVERT THEM TO .PDF
#THE FUNCTION CREATES A TEMPORAL FOLDER TO KEEP ALL THE DOCUMENTS IN .PDF
#OTHER APPROACH WOULD BE READ ALL THE DOCUMENTS IN .PDF, THUS CONVERT THEM ALL OF THEM, HOWEVER DOES NOT GURANTEE 
#APPROPIATE READING FOR SOME INFORMATION
def get_specifications(PATH_R_DOCS_2,j):
    #os.mkdir('temp2') #CREATES THE FOLDER
  
    data_specifications_dose = []
    number_specification_fractions = []
    #dir_list = os.listdir(PATH_R_DOCS_2)

    #for j in dir_list:    
    generate_pdf(PATH_R_DOCS_2+j, 'temp2') 
  
    pdf_reader = PdfReader(open('temp2/'+j.split('.')[0]+'.pdf', "rb"))

    page = pdf_reader.pages[0]
      
    for p in str(page.extract_text()).lower().split('\n'):
        if 'total dose:' in p:
            if 'specify:' in p:
                text = p.replace('\u2002\u2002\u2002\u2002\u2002 ','')#ERASES SOME EXTRA INFO/SYMBOLS NOT WANTED
                text2 = text.split('specify:')[-1].strip()
                  
                data_specifications_dose.append(text2)
                  
        if 'number of fractions remaining:' in p:
            if 'specify:' in p:
                  
                text = p.replace('\u2002\u2002\u2002\u2002\u2002 ','')
                text2 = text.split('specify:')[-1].strip()                  
                number_specification_fractions.append(text2) 
                
    os.remove("temp2/"+j.split('.')[0]+'.pdf')
    #os.remove("temp2")
    return data_specifications_dose[0],number_specification_fractions[0]
  
#SAVE THE DATA IN .CSV FILE 
#TO DO: SAVE IT IN .JSON
def save_docsv2(ids_mrn,data):
    data_w_id = []
    for j in ids_mrn:
        for k in data:
            if str(j[1]) in k[0]:
                  data_w_id.append([j[0],k[0],k[1],k[2]])
    comment_data = pd.DataFrame(data_w_id ,columns=['id','mrn','comment'])
    comment_data.to_csv('OUTPUT_DOC.csv',index=False)
  
    return
  
def save_docs(ids_mrn,data):
    data_w_id = []
    for j in ids_mrn:
        for k in data:
            if str(j[1]) in k[0]:
                data_w_id.append([j[0],k[0],k[1]])
    comment_data = pd.DataFrame(data_w_id ,columns=['id','mrn','comment'])
    comment_data.to_csv('OUTPUT_DOC.csv',index=False)
  
    return 
    
###################################################################################################3

###CHANGE HERE THE PATHS 
PATH_DOCS = '/data2/odette/aria_RADO_DOCS/RADO_ONCS/'
PATH_R_DOCS = '/mnt/iDriveShare/Yossera_HNC_Summer/Replan_new_documents/'
PATH_DOCS_N = '/data2/odette/aria_RADO_DOCS/NUT_ONCS/'

#dir_list = os.listdir(PATH_DOCS)
#dir_list_R = os.listdir(PATH_R_DOCS)
#dir_list_N = os.listdir(PATH_DOCS_N)

data_ids_mrn = pd.read_csv('/mnt/iDriveShare/Kayla/for_yossera_new_patient_list.csv')

#GETS THE MRN FOR EACH PATIENT ID
ids_mrn = []
for j in data_ids_mrn[['id','mrn']].values:
    if list(j) not in ids_mrn:
          ids_mrn.append(list(j))
    


dir_list = os.listdir(PATH_DOCS_R)
data_all = []

if os.path.isdir('temp2')==False:
    os.mkdir('temp2')
        
if os.path.isdir('temp')==False:
    os.mkdir('temp')
        
        
for j in dir_list:
    mrn = j.split('.')[0]
    specifications_dose,specifications_fx = get_specifications(PATH_DOCS_R,j)
    dates = replan_date(PATH_DOCS_R,j)
    signature = get_signature(PATH_DOCS_R,j)   
    
    comments_instructions = read_comments_instructions_replan_docs(PATH_DOCS_R,j)
    total_dose,fractions_remaining = read_fractions_remaining_and_total_dose(PATH_DOCS_R,j)
    
    comments_physics = get_comments_in_physics_section(PATH_DOCS_R,j)
    
    reasons,specifications0 = read_checkboxes_reason(PATH_DOCS_R,j)
    who_noticed_during = read_checkboxes_who_when_noticed(PATH_DOCS_R,j)
    recommendation_replan_data = read_checkboxes_recommendation(PATH_DOCS_R,j)
    
    if signature=='':
        if comments_physics=='':
            discussion = 'NO'
        else:
            discussion = 'YES'
    else:
        discussion = 'YES'
        
    for ids in ids_mrn:
        if str(ids[1]) in str(mrn):
            #print(ids[0])
            data_all.append([int(ids[0]),str(ids[1]),dates,signature,recommendation_replan_data[0],recommendation_replan_data[1],total_dose,specifications_dose,fractions_remaining,specifications_fx,
                        reasons,specifications0,who_noticed_during[0],who_noticed_during[1],
                        who_noticed_during[2],who_noticed_during[3],comments_instructions,discussion,comments_physics])
    
data_all.sort(key=lambda data_all: data_all[0])

df = pd.DataFrame(data_all,columns=['id','mrn','date','signed by','recommendation','replan or not','total dose','specification in dose',
                                  'fxs remaining','specification in fxs','reason for replanning',
                                  'reason specifications','noticed by','details noted','noticed during',
                                  '(during) details','comments/instructions','discussed with physics',
                                  'physics comments'])
                                  
df.to_json(orient='records')

result = df.to_json(orient="index")
parsed = loads(result)
json1 = dumps(parsed, indent=4)

with open('replan_docs_data.json', 'w') as outfile:
    outfile.write(json1)


